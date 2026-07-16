"""
生成符合《智能系统学报》格式的完整 Word 论文文档
含全部4张配图、标准期刊排版
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = OUTPUT_DIR  # figures are here
DOC_PATH = os.path.join(os.path.dirname(OUTPUT_DIR), 'SCOPE-学术论文-智能系统学报.docx')

doc = Document()

# ============================================================
# Page setup: A4, standard margins
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.size = Pt(10.5)  # 五号
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进

# ============================================================
# Helper functions
# ============================================================
def add_paragraph(text, font_name='宋体', font_size=Pt(10.5), bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
                  space_before=Pt(0), space_after=Pt(0), color=None, italic=False):
    """Add a paragraph with specified formatting"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_heading_custom(text, font_name='黑体', font_size=Pt(14), bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(12), space_after=Pt(6)):
    """Add a section heading"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = space_before
    pf.space_after = space_after
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p

def add_figure(image_path, caption_cn, caption_en, width_inches=5.8):
    """Add a figure with bilingual caption"""
    if not os.path.exists(image_path):
        print(f"  WARNING: Figure not found: {image_path}")
        return

    # Add the image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(3)
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(width_inches))

    # Chinese caption
    add_paragraph(caption_cn, font_name='黑体', font_size=Pt(9), bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
                  space_before=Pt(2), space_after=Pt(0))

    # English caption
    add_paragraph(caption_en, font_name='Times New Roman', font_size=Pt(9), bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
                  space_before=Pt(0), space_after=Pt(6), italic=True)

def add_subheading(text, level=2):
    """Add subsection heading: level 2=黑体5号, level 3=楷体5号"""
    if level == 2:
        font_name, font_size = '黑体', Pt(10.5)
    else:
        font_name, font_size = '楷体', Pt(10.5)
    add_heading_custom(text, font_name=font_name, font_size=font_size,
                        space_before=Pt(8), space_after=Pt(4))

print("Generating Word document...")

# ============================================================
# TITLE (小2号黑体 = 18pt)
# ============================================================
add_paragraph('SCOPE：一种面向AI应用场景发现的系统性方法论',
              font_name='黑体', font_size=Pt(18), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(24), space_after=Pt(8))

# ============================================================
# AUTHORS (小4号楷体 = 12pt)
# ============================================================
add_paragraph('罗毅晗',
              font_name='楷体', font_size=Pt(12), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(4), space_after=Pt(4))

# ============================================================
# AFFILIATION (小5号楷体 = 9pt)
# ============================================================
add_paragraph('（中国航发湖南动力机械研究所，湖南 株洲 412002）',
              font_name='楷体', font_size=Pt(9), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(0), space_after=Pt(12))

# ============================================================
# CHINESE ABSTRACT
# ============================================================
add_paragraph('摘  要', font_name='黑体', font_size=Pt(9), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(6), space_after=Pt(2))

abstract_cn = (
    '如何系统性地发现和规划人工智能（AI）在业务中的应用场景，是当前AI产业化落地的核心瓶颈。'
    '现有方法多依赖个人经验和直觉进行"拍脑袋"式的场景挖掘，缺乏可重复、可验证的系统性方法论。'
    '本文借鉴TRIZ理论的思想结构（矛盾分析→发明原理→进化法则→效应库），提出了一种面向AI应用场景发现的'
    '系统性方法论——SCOPE。该方法论包含五个层次：第一层定义6个业务改善参数与6个恶化约束参数，构成矛盾建模语言；'
    '第二层构建了包含36个AI业务发明原理的五大类原理库，每个原理均经过真实案例验证并配备标准化模板；'
    '第三层建立6×6矛盾矩阵，通过210个真实案例的交叉统计，为每个矛盾组合推荐最优原理，并标注信度等级与禁忌标记；'
    '第四层定义7条AI业务进化路线，为场景规划提供时间维度；第五层构建AI能力效应库，实现从"发明原理"到"具体AI技术"的映射。'
    '在此基础上，开发并部署了SCOPE Agent交互式系统，将五步法流程产品化。'
    '210个案例的统计分析表明，认知自动化（63次验证）与预测推演（60次验证）构成AI赋能业务的两大基础能力。'
    '矛盾矩阵填满率达100%，其中运营效率×实施复杂度（N≈55）为最强矛盾-原理组合。'
    '该方法论将AI应用场景的发现从依赖个人灵感的"艺术"转变为可操作、可复现的"科学"，'
    '为产学研各界提供了系统性的AI场景规划工具。'
)
add_paragraph(abstract_cn, font_name='宋体', font_size=Pt(9),
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(0), space_after=Pt(4))

# Keywords
kw_cn = ('关键词：人工智能应用；场景发现；TRIZ；矛盾矩阵；发明原理；方法论；AI赋能；业务创新；进化路线；效应库')
add_paragraph(kw_cn, font_name='黑体', font_size=Pt(9), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(0), space_after=Pt(2))

# CLC number
add_paragraph('中图分类号：TP301', font_name='宋体', font_size=Pt(9),
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(0), space_after=Pt(10))

# ============================================================
# ENGLISH TITLE (小3号 Times New Roman Bold = 15pt)
# ============================================================
add_paragraph('SCOPE: A Systematic Methodology for AI Application Scenario Discovery',
              font_name='Times New Roman', font_size=Pt(15), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(16), space_after=Pt(6))

# English Authors
add_paragraph('LUO Yihan',
              font_name='Times New Roman', font_size=Pt(12), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(2), space_after=Pt(2))

# English Affiliation
add_paragraph('(AECC Hunan Aviation Powerplant Research Institute, Zhuzhou 412002, China)',
              font_name='Times New Roman', font_size=Pt(9), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
              space_before=Pt(0), space_after=Pt(10))

# English Abstract
add_paragraph('Abstract', font_name='Times New Roman', font_size=Pt(9), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(4), space_after=Pt(2))

abstract_en = (
    'How to systematically discover and plan artificial intelligence (AI) application scenarios in business '
    'remains a core bottleneck in AI industrialization. Current approaches predominantly rely on individual '
    'experience and intuition for ad-hoc scenario mining, lacking a repeatable and verifiable systematic methodology. '
    'Drawing inspiration from the structural framework of TRIZ theory (contradiction analysis → inventive principles '
    '→ evolution laws → effect database), this paper proposes SCOPE, a systematic methodology for AI application '
    'scenario discovery. The methodology comprises five layers: Layer 1 defines 6 business improvement parameters '
    'and 6 constraint deterioration parameters, forming a contradiction modeling language; Layer 2 constructs a '
    'five-category library of 36 AI business inventive principles, each validated by real-world cases and equipped '
    'with standardized templates; Layer 3 establishes a 6×6 contradiction matrix, recommending optimal principles '
    'for each contradiction combination through cross-statistical analysis of 210 real cases, annotated with '
    'confidence levels and contraindication markers; Layer 4 defines 7 AI business evolution trajectories, '
    'providing a temporal dimension for scenario planning; Layer 5 builds an AI capability effect database, '
    'mapping from inventive principles to specific AI technologies. An interactive SCOPE Agent system was '
    'developed and deployed to productize the five-step workflow. Statistical analysis of 210 cases demonstrates '
    'that cognitive automation (63 validations) and predictive simulation (60 validations) constitute the two '
    'foundational capabilities of AI-enabled business. The contradiction matrix achieves 100% fill rate. '
    'This methodology transforms AI application scenario discovery from an intuition-dependent "art" into '
    'an operable and reproducible "science," providing a systematic AI scenario planning tool for academia and industry.'
)
add_paragraph(abstract_en, font_name='Times New Roman', font_size=Pt(9),
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(0), space_after=Pt(4))

kw_en = ('Keywords: artificial intelligence application; scenario discovery; TRIZ; contradiction matrix; '
         'inventive principles; methodology; AI empowerment; business innovation; evolution trajectory; effect database')
add_paragraph(kw_en, font_name='Times New Roman', font_size=Pt(9), bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0.74),
              space_before=Pt(0), space_after=Pt(10))

# ============================================================
# 0. INTRODUCTION (引言 — 不使用标题)
# ============================================================
intro_text = (
    '人工智能（AI）技术正经历从"技术探索"到"产业落地"的关键转型期。以GPT-4、Claude、DeepSeek为代表的大语言模型'
    '在自然语言理解、多模态感知、逻辑推理等能力上取得了突破性进展[1-3]，AI赋能各行各业的潜力已得到广泛认可。'
    '然而，产业界面临的核心挑战已从"AI能做什么"转变为"AI应该用在哪里、怎么用"——即AI应用场景的发现与规划问题[4-5]。'
)
add_paragraph(intro_text, font_name='宋体', font_size=Pt(10.5),
              space_before=Pt(0), space_after=Pt(2))

intro2 = (
    '当前的AI应用场景发现主要依赖三种路径：一是"技术驱动"路径，由AI技术提供方主导，从技术能力出发寻找适配场景，'
    '容易落入"拿着锤子找钉子"的陷阱[6]；二是"问题驱动"路径，由业务方主导，从痛点出发寻求技术方案，'
    '但常受限于对AI能力的认知边界[7]；三是"模仿跟随"路径，通过观察行业标杆或竞争对手的AI实践进行复制，'
    '但缺乏对自身业务独特性的深入分析[8]。这三种路径的共性问题在于：缺乏一套可重复、可验证的系统性方法论，'
    '场景发现的质量高度依赖于个人经验和灵感——即"拍脑袋"。'
)
add_paragraph(intro2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

intro3 = (
    '这一困境与20世纪中期工程技术领域的状态高度相似。当时，技术创新同样依赖个人天才和随机灵感，'
    '直到前苏联发明家根里奇·阿奇舒勒（Genrich Altshuller）提出了TRIZ理论（"发明问题解决理论"的俄语缩写），'
    '通过对4万份高水平专利的统计分析，提炼出39个通用工程参数、40个发明原理和矛盾矩阵，'
    '将技术创新从一个"艺术"转变为一门"科学"[9-11]。TRIZ的核心洞察在于：'
    '不同领域的技术问题背后，存在着共通的矛盾模式；识别这些模式，就能系统性地生成创新解决方案。'
)
add_paragraph(intro3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

intro4 = (
    '本文的核心假设是：AI应用场景的发现也遵循类似的矛盾模式——不同行业的业务场景背后，'
    '存在着共通的"想要提升什么（改善）vs 担心什么恶化（约束）"的矛盾结构；识别这些矛盾，'
    '就能系统性地推荐适合的AI赋能策略。基于这一假设，本文借鉴TRIZ的思想结构，但基于210个真实AI应用案例的统计分析，'
    '构建了一套完全独立的、面向AI应用场景发现的方法论——SCOPE。'
)
add_paragraph(intro4, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

intro5 = (
    'SCOPE的名称来自其五步法的首字母缩写：Strengthen（增强目标）→ Constrain（约束预警）→ '
    'Originate（创意生成）→ Path（路径规划）→ Effect（效应匹配）。同时，"scope"一词在英文中意为"范围、领域"，'
    '契合方法论的核心目标——构建AI业务场景机会的全景图。'
)
add_paragraph(intro5, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

intro6 = (
    '本文的主要贡献包括：（1）提出了SCOPE方法论的整体框架，定义了五步法工作流程；'
    '（2）从210个真实案例中自下而上地提炼了6+6个业务场景参数，建立了矛盾建模语言；'
    '（3）构建了包含36个AI业务发明原理的五大类原理库，每个原理配有标准化五段模板；'
    '（4）建立了6×6矛盾矩阵，通过案例交叉统计实现了100%填满率，并标注了信度等级与禁忌组合；'
    '（5）定义了7条AI业务进化路线和AI能力效应库，形成完整的五层架构；'
    '（6）开发并部署了SCOPE Agent交互式系统，实现了方法论的工程化应用。'
)
add_paragraph(intro6, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 1. RELATED WORK
# ============================================================
add_heading_custom('1  相关工作')

add_subheading('1.1  TRIZ理论概述')
r1_1 = (
    'TRIZ是由前苏联发明家阿奇舒勒于1946年创立的系统性创新方法论[9]。其核心发现是：不同领域的发明问题背后'
    '存在共通的矛盾模式，而解决这些矛盾的方法也是可归纳、可复用的。TRIZ的主要工具体系包括：'
    '39个通用工程参数（如"运动物体的重量""速度""力""可靠性"等），用于将具体技术问题抽象为标准矛盾——'
    '"我想改善参数A，但这会导致参数B恶化"；40个发明原理（如"分割""抽取""局部质量""嵌套"等），每个原理包含若干子原理，'
    '是被大量高水平专利验证过的通用创新策略；矛盾矩阵（39×39），横轴为恶化参数，纵轴为改善参数，交叉格中填入最常用的'
    '发明原理编号；技术进化法则，描述技术系统从简单到复杂、从宏观到微观的进化规律；科学效应库，将具体的物理/化学/几何效应'
    '与工程需求配对。TRIZ经过70余年的发展，已被全球大量企业（三星、波音、宝洁、华为等）应用于产品创新和工艺改进[12-14]。'
    '檀润华等学者对TRIZ在中国的推广和应用做出了重要贡献[15]。然而，TRIZ的设计初衷是解决工程技术领域的物理/机械创新问题，'
    '其39个工程参数（如"应力""温度""能量损失"）与AI应用场景的语言鸿沟巨大，直接套用并不现实。'
)
add_paragraph(r1_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('1.2  AI应用场景发现方法')
r1_2_1 = (
    '近年来，学术界和产业界开始关注AI应用场景的系统化发现问题。在学术层面，Barnett等[16]对计算研究中场景方法的使用'
    '进行了系统综述（59篇文献），发现多数场景构建仍以人工为主，缺乏AI增强的系统工具。Negrini[17]对生成式AI用于场景生成'
    '的文献进行了系统综述，发现该方向仍处于早期发展阶段（20篇主要文献，2013-2024）。PRISE框架[18]整合了设计思维、'
    '精益创业与工程实践，提供了从概念到实施的AI产品孵化流程，但侧重于产品开发流程管理而非场景的创造性发现。'
    'IPAPS框架[19]针对中小企业提出了AI产品与服务的创新流程，但同样以流程管理为主。'
)
add_paragraph(r1_2_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

r1_2_2 = (
    '在产业层面，Merkle于2025年提出的AI商业价值画布（AIBC）[20]将设计思维方法应用于AI规划，通过九大构建模块帮助企业'
    '识别AI价值机会。Scheichl[21]提出了AI增强的双钻石模型，将AI技术映射到"发现→定义→开发→交付"全流程。'
    '然而，这些方法本质上仍是将AI视为"设计工具的增强组件"而非"独立的方法论内核"。'
    '中文文献方面，赵健等[22]对人工智能大模型及其应用进行了全面综述，但侧重于技术能力描述而非场景发现方法。'
    '大模型行业应用设计与实践[23]等研究提供了行业级实践案例，但未形成可推广的通用方法论。'
)
add_paragraph(r1_2_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('1.3  现有方法的局限性')
r1_3 = (
    '综合以上分析，现有方法存在以下核心局限：（1）缺乏矛盾建模语言——大多数方法直接跳入"AI能解决什么问题"的讨论，'
    '缺少"业务目标 vs 约束条件"的结构化矛盾建模环节，容易在实施后引发意外的负面后果。'
    '（2）缺乏基于实证统计的推荐体系——现有的"AI应用场景匹配"建议通常来自专家经验或启发式规则，缺乏基于大规模案例统计的、'
    '可量化信度的推荐体系。（3）缺乏时间维度的进化指引——AI能力处于快速演进中，当前的"最优方案"可能在18个月后过时。'
    '（4）缺乏从原理到技术的映射机制——业务人员即使知道"应该用AI做个性化推荐"，也不清楚具体需要哪些AI技术组件。'
    'SCOPE方法论的提出，正是为了填补以上四个空白。'
)
add_paragraph(r1_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 2. SCOPE OVERALL FRAMEWORK
# ============================================================
add_heading_custom('2  SCOPE方法论总体框架')

add_subheading('2.1  五步法概述')
s2_1 = (
    'SCOPE方法论的核心是一个五步迭代流程，每一步对应方法论名称的一个字母。步骤1—Strengthen（增强目标）：'
    '从6个业务改善参数中，明确要增强的核心业务目标，关键问题是"我最想提升的是什么？"。'
    '步骤2—Constrain（约束预警）：从6个恶化约束参数中，识别引入AI后最担心的副作用，'
    '关键问题是"引入AI后，我最怕什么出问题？"。步骤3—Originate（创意生成）：将步骤1和2确定的矛盾对'
    '输入矛盾矩阵，获得按信度排序的AI业务发明原理推荐。步骤4—Path（路径规划）：评估当前业务场景在7条进化路线上的位置，'
    '规划从"近期（复制/自动化）→ 中期（增强/预测）→ 远期（自主/创造）"的三级跳路径。'
    '步骤5—Effect（效应匹配）：根据推荐的发明原理，查效应库找到匹配的具体AI技术（如LLM、计算机视觉、时序预测等），'
    '形成完整的技术选型方案。这五个步骤构成了从"模糊的业务痛点"到"清晰的AI应用场景方案"的完整闭环。'
)
add_paragraph(s2_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('2.2  五层架构')
s2_2 = (
    '与五步法相对应，SCOPE建立了五个相互关联的知识层：第一层为业务场景参数（6个改善参数+6个恶化参数+物理矛盾），'
    '第二层为AI业务发明原理库（36个原理，五大类，标准化五段模板），第三层为矛盾矩阵'
    '（6×6=36格，210案例交叉统计，信度+禁忌标记，100%填满率），第四层为AI业务进化路线'
    '（7条进化路线+回滚分支），第五层为AI能力效应库（"业务需求→AI技术"配对映射，共14项映射）。'
    '这些层次由底向上逐层支撑：参数层提供矛盾建模的语言，原理层提供解决方案的候选空间，'
    '矩阵层提供"矛盾→原理"的推荐引擎，进化路线层提供时间维度的推进策略，效应库层提供"原理→技术"的实现桥梁。'
)
add_paragraph(s2_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

# FIGURE 1: Architecture
add_figure(
    os.path.join(FIG_DIR, 'fig1_scope_architecture.png'),
    '图1  SCOPE方法论五层架构与五步法流程',
    'Fig.1  Five-Layer Architecture and Five-Step Workflow of the SCOPE Methodology',
    width_inches=5.5
)

add_subheading('2.3  与TRIZ的关系')
s2_3 = (
    '需要强调的是：SCOPE是完全独立的方法论。它与TRIZ的关系是"结构借鉴"而非"内容移植"。具体而言，'
    'TRIZ的39个工程参数基于4万份专利，SCOPE的6+6个业务参数基于210个AI案例独立定义；TRIZ的40个工程原理与SCOPE的36个AI业务'
    '原理各自独立提炼；矛盾矩阵方面TRIZ为39×39而SCOPE为6×6；进化法则方面TRIZ有8条技术进化法则而SCOPE定义7条业务进化路线；'
    '效应库方面TRIZ为物理/化学/几何效应数据库而SCOPE为AI技术能力映射表。SCOPE的独有特征包括：基础原理标记、禁忌标记、'
    '信度等级、五段标准化模板和三级跳路径规划。这种关系类似于：TRIZ为"如何系统性地发现解决方案"提供了一个思想原型，'
    'SCOPE将该原型的骨架结构应用于一个全新的领域——AI业务应用场景发现，并在此过程中重新定义了所有内容。'
)
add_paragraph(s2_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 3. BUSINESS SCENARIO PARAMETERS
# ============================================================
add_heading_custom('3  业务场景参数体系')

s3_intro = (
    'SCOPE的第一层为矛盾建模提供了统一的参数语言。参数的定义遵循"自下而上涌现"原则——从案例的语义聚类中自然形成类别边界，'
    '而非自上而下地预设分类框架。经80个案例时的初步聚类，到150案例时得到确认，最终在210案例规模下稳定为6个改善参数与6个恶化参数。'
)
add_paragraph(s3_intro, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('3.1  改善参数（Strengthen）')
s3_1 = (
    '改善参数回答"你想提升什么？"这一问题，包括：S1运营效率（流程吞吐速度、人力替代、资产利用率、成本控制，210案例覆盖率73%），'
    'S2客户满意度（用户体验、NPS评分、信任度、参与度，覆盖率31%），S3创新能力（新产品/服务/模式的创造——"之前做不到的事现在能做了"，'
    '覆盖率33%），S4风险控制（安全、合规、欺诈检测、故障预防、质量保障，覆盖率28%），S5收入增长（GMV、利润率、新客户获取、'
    '客单价提升，覆盖率21%），S6决策速度（从信息到决策的周期缩短——人的决策质量和速度，覆盖率16%）。'
    'S1（运营效率）以73%的覆盖率遥遥领先，这与当前的产业现状高度吻合——AI在效率和自动化方面的应用最为成熟。'
    'S6（决策速度）覆盖率最低（16%），反映出"用AI替代或增强人类决策"仍处于探索早期。'
)
add_paragraph(s3_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('3.2  恶化参数（Constrain）')
s3_2 = (
    '恶化参数回答"引入AI后，你最担心什么？"这一问题，包括：C1实施复杂度（技术可行但不一定能落地——集成难度、人才缺口、'
    '组织变革阻力，覆盖率50%），C2可解释性（AI决策的黑箱性——监管要求、用户信任、错误追溯，覆盖率30%），'
    'C3对人的依赖度（AI越强人越弱——技能退化、组织抗拒、人性化服务缺失，覆盖率24%），'
    'C4成本（初始投资 vs 长期回报的张力——算力、数据、人才、维护，覆盖率28%），'
    'C5数据隐私风险（用户数据、商业机密、跨境合规，覆盖率16%），'
    'C6系统安全性（AI出错的直接后果——金融损失、安全事故、公共安全威胁，覆盖率18%）。'
    'C1（实施复杂度）以50%覆盖率成为最普遍的约束，印证了一个行业现实：AI落地的最大障碍往往不是技术可行性，而是组织、人才和变革管理的复杂度。'
)
add_paragraph(s3_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('3.3  物理矛盾')
s3_3 = (
    '除了技术矛盾（改善参数×恶化参数），SCOPE还识别了一类特殊的矛盾形态——物理矛盾：一个业务要素同时面临对立要求。'
    '例如："客服既要7×24小时高响应（效率），又要有温度、能解决复杂问题（质量）"；'
    '"库存既要少以降低成本（成本），又要保证永不断货（服务）"；'
    '"AI既要自主决策提升效率，又要每个决策可追溯可解释（可解释性）"；'
    '"个性化推荐既要精准，又不能侵犯用户隐私（数据隐私）"。'
    '物理矛盾由分离原理（时间分离、空间分离、条件分离、系统级别分离）解决，是SCOPE方法论中最精妙也最具挑战性的矛盾类型。'
)
add_paragraph(s3_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 4. AI BUSINESS INVENTIVE PRINCIPLES
# ============================================================
add_heading_custom('4  AI业务发明原理库')

s4_intro = (
    'SCOPE的第二层是包含36个AI业务发明原理的五大类原理库。每个原理必须满足以下质量标准：'
    '（1）在至少2个独立案例中得到验证；（2）通过"定义→核心机制→适用信号→典型案例→注意事项"五段模板进行标准化描述；'
    '（3）按信度等级和特殊标签进行分类。'
)
add_paragraph(s4_intro, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('4.1  原理库设计原则')
s4_1 = (
    '原理库的设计遵循四项核心原则：自下而上涌现——原理从210个案例中迭代聚类形成，非预先设定，原始V0版本的40个原理经'
    '80案例反向验证，执行4项合并，精炼为36个；实证驱动——每个原理的验证次数来自真实案例统计，基础原理要求在矛盾矩阵中'
    '默认推荐；透明信号——每个原理配备适用信号（何时应该考虑用）、禁忌信号（特定组合下有失败风险）和区分信号（与相似原理的对比说明）；'
    '动态演化——随着案例增长，原理可以合并、拆分或新设，编号空白（原理32退役）保留以追踪变更历史。'
)
add_paragraph(s4_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('4.2  五大类原理结构与三大支柱')
s4_2 = (
    '36个原理分为五大类：第一大类"重构流程与体验"（5个原理），聚焦于重塑业务流程和用户体验；'
    '第二大类"增强决策与洞察"（5个原理），赋予系统或人更强的认知和决策能力；'
    '第三大类"重塑生产关系"（10个原理），改变人、机器、组织的协作方式；'
    '第四大类"颠覆产品与资产"（6个原理），重新定义"产品"和"资产"的形态；'
    '第五大类"范式跃迁"（10个原理），颠覆游戏规则本身。'
)
add_paragraph(s4_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s4_2b = (
    '210案例统计中，三个原理的验证次数遥遥领先，构成SCOPE的"三大支柱"：'
    '原理2"认知自动化"（63次验证，定义为"AI看懂、听懂、读懂非结构化信息"）、'
    '原理7"预测与推演"（60次验证，定义为"从事后分析到事前预判"）、'
    '原理9"智能编排与调度"（43次验证，定义为"在复杂约束下动态制定最优行动计划"）。'
    '原理2和原理7已被标记为基础原理——在矛盾矩阵中默认推荐且不参与排序，腾出空间给差异化原理。'
    '它们共同构成了AI赋能业务的"双引擎"：认知自动化解决"信息理解"问题，预测与推演解决"决策前瞻"问题。'
)
add_paragraph(s4_2b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('4.3  原理模板示例：智能分割')
s4_3 = (
    '以原理1"智能分割（Intelligent Task Segmentation）"为例，展示标准化五段模板。'
    '定义：将原本由人统一处理的复杂任务，拆解后让AI处理标准、高并发部分，人专注高价值例外。'
    '核心机制：任务的可分割性是效率提升的前提，AI擅长处理"长尾中的头部"——大量重复、边界清晰的工作；'
    '人擅长处理"头部中的长尾"——少量复杂、需要综合判断的例外。'
    '适用信号：业务流程中有明确的"标准操作"和"例外处理"两层，标准操作占比>70%，例外处理的判断标准可以被文档化。'
    '典型案例：客服拆分为AI处理常见问题，人工只介入复杂投诉（平安集团）；海底捞AI做标准化工作（数据分析/库存预测/排班），'
    '人做情感沟通与个性化服务；HR招聘拆分为AI初筛+初面，HR专注终面和决策（智联招聘）。'
    '注意事项：分割边界需要动态调整——AI能力的提升会持续改变最优人机边界；注意分割后的"交接成本"不应吞噬效率收益。'
    '全部36个原理均采用此模板进行标准化描述（完整模板见SCOPE方法论文档）。'
)
add_paragraph(s4_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 5. CONTRADICTION MATRIX
# ============================================================
add_heading_custom('5  矛盾矩阵')

s5_intro = (
    '矛盾矩阵是SCOPE的核心工具，也是方法论从"定性建议"升维到"统计推荐"的关键。它将步骤1（改善目标）和步骤2（恶化约束）'
    '形成的矛盾对，映射到步骤3（AI发明原理推荐），实现"给定矛盾，返回最优原理"的查询功能。'
)
add_paragraph(s5_intro, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('5.1  矩阵构建方法与信度体系')
s5_1 = (
    '矛盾矩阵为6×6结构：横轴为6个恶化参数（C1-C6），纵轴为6个改善参数（S1-S6），共36个单元格。'
    '每个单元格的内容基于210个真实AI应用案例的交叉标注统计生成——对每一个案例，标注其改善目标参数和恶化约束参数，'
    '然后统计该"矛盾对"下各类原理的出现频次。单元格内的原理按出现频次降序排列。基础原理（🔧2和🔧7）以特殊标注形式'
    '列在列表末尾，不参与频次排序，体现"默认可用但不干扰推荐"的设计意图。每个单元格附有信度等级：'
    '🟢高信度（N≥15）、🟡中信度（N=6-14）、🟠低信度（N=3-5）、🔴最低信度（N=1-2）。'
)
add_paragraph(s5_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

# FIGURE 2: Contradiction Matrix
add_figure(
    os.path.join(FIG_DIR, 'fig2_contradiction_matrix.png'),
    '图2  SCOPE矛盾矩阵——6×6=36格，210案例交叉统计',
    'Fig.2  SCOPE Contradiction Matrix — 6×6 cells, 210 cases cross-statistics',
    width_inches=5.5
)

add_subheading('5.2  关键发现与禁忌组合')
s5_2 = (
    '矛盾矩阵分析揭示了若干关键发现。最强矛盾-原理组合为S1（运营效率）× C1（实施复杂度），以N≈55的案例量位居第一，'
    '推荐原理包括智能分割、智能编排、数字孪生、服务机器人化等9个差异化原理——这是AI产业化的主战场，'
    '在控制实施复杂度的前提下提升运营效率，是大多数AI应用的基本命题。第二强组合为S3（创新能力）× C1（实施复杂度），'
    '以N≈24的案例量位居第二，推荐原理高达14个，覆盖面最广——反映了"用AI创造新产品/新服务"的路径多样性。'
    '在210案例规模下，矛盾矩阵36格全部有数据（填满率100%），其中高信度7格（19.4%），中信度12格（33.3%），'
    '低信度8格（22.2%），最低信度9格（25.0%）。'
)
add_paragraph(s5_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s5_2b = (
    '一个重要发现是S5（收入增长）× C2（可解释性）矛盾对存在禁忌组合——携程AI调价助手因价格算法不透明（可解释性恶化），'
    '于2026年3月被指控"大数据杀熟"而被迫下线。这是210案例库中最重要的失败教训之一。该格的推荐原理列表中标有"原理4（慎）"，'
    '提示用户：若必须使用原理4（动态定价与匹配），必须附带公平性与透明度保障机制。'
)
add_paragraph(s5_2b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 6. EVOLUTION TRAJECTORIES & EFFECT DATABASE
# ============================================================
add_heading_custom('6  进化路线与效应库')

add_subheading('6.1  七条AI业务进化路线')
s6_1 = (
    'SCOPE定义了7条描述AI业务应用从低阶到高阶的进化路线。路线1"洞察深度"：报表（发生了什么）→ 仪表盘（正在发生什么）→ '
    '预测性警报（将要发生什么）→ 规范建议（该怎么办）→ 自主行动。路线2"决策自主度"：AI给出建议 → AI和人协同决策 → '
    'AI增强人 → AI替代人（特定任务）→ AI自主执行、人监督 → AI完全自主。路线3"产品形态"：流程自动化 → 智能产品化 → '
    '生态平台化。路线4"技术范式"：基于规则 → 基于模型 → 基于生成 → 基于自进化。路线5"交互方式"：菜单/表单 → 自然语言 → '
    '多模态融合 → 无感交互（环境感知）。路线6"Agent架构"：单Agent → 多Agent协同 → Agent生态（Agent间自主协作），'
    '这是2025-2026年最活跃的趋势方向。路线7"价值分配"：层级中介 → 平台中介 → AI使能的去中介化。'
    '企业可以根据自身业务场景，评估在每条路线上的当前位置并规划下一站。'
)
add_paragraph(s6_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

# FIGURE 4: Evolution Routes
add_figure(
    os.path.join(FIG_DIR, 'fig4_evolution_routes.png'),
    '图4  SCOPE七条AI业务进化路线',
    'Fig.4  Seven AI Business Evolution Trajectories of SCOPE',
    width_inches=5.5
)

add_subheading('6.2  回滚分支与AI能力效应库')
s6_2 = (
    'SCOPE的一个重要创新是引入了"回滚"概念——进化不是单向的。当恶化参数失控时，AI应用可能从高阶阶段回退到低阶阶段。'
    '回滚的触发条件包括：信任危机（客户满意度急剧下降）、监管介入、安全事故、以及ROI持续为负。'
    '210案例库中收录了携程AI调价助手从"AI动态定价"回退到"固定定价"的回滚案例，为回滚机制提供了实证依据。'
)
add_paragraph(s6_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s6_2b = (
    'SCOPE的第五层建立了从"业务需求"到"具体AI技术"的效应映射。与TRIZ的科学效应库概念相似，当用户从矛盾矩阵获得了推荐的'
    '发明原理，效应库提供实现该原理所需的具体AI技术选型建议。当前效应库包含14项映射，覆盖了主要的AI技术领域：'
    '大语言模型（LLM）用于理解/生成人类语言；计算机视觉（CNN/ViT/YOLO）用于图像/视频对象识别；'
    'OCR+信息抽取+LLM用于文档结构化信息提取；时序预测（LSTM/Transformer/GBDT）用于数值趋势预测；'
    '孤立森林/自编码器/GNN用于异常/欺诈模式发现；协同过滤/上下文Bandit/LLM用于个性化推荐；'
    '图嵌入/路径排序/LLM+KG用于知识图谱推理与溯源；强化学习/运筹优化/约束求解用于多约束动态调度优化；'
    '行为基线建模+自主决策引擎用于自主安全威胁响应；Neural CAD/扩散模型/物理仿真用于3D几何/结构生成；'
    '分布式MPC/注意力通信/集中式调度用于多智能体蜂群协调；区块链+Oracle+ML验证用于自动化协议执行；'
    '多LLM+推理聚合+统计校准用于多模型群体审议；TTS+口型生成+LLM对话用于数字人/声线克隆。'
)
add_paragraph(s6_2b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 7. EMPIRICAL VALIDATION
# ============================================================
add_heading_custom('7  实证验证：210案例分析')

add_subheading('7.1  案例收集与标注方法')
s7_1 = (
    'SCOPE的实证基础来自210个真实AI应用案例的系统收集、标注与统计分析。案例来源包括：公开报道的AI应用成功案例（占比约83%），'
    '学术论文中描述的AI应用实例（占比约10%），以及行业报告和白皮书中的案例（占比约7%）。每个案例需满足以下收录标准：'
    '具备明确的业务应用场景描述、可识别的AI技术方案、以及可验证的业务效果指标。案例标注采用人工标注+交叉验证的方式，'
    '每个案例标注改善参数、恶化参数、适用的发明原理、所属行业、地域分布、时间信息和成功/失败状态。'
)
add_paragraph(s7_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('7.2  案例覆盖概况与频次排名')
s7_2 = (
    '210案例覆盖42+行业领域，主要包括：金融/保险、制造/工业、医疗/健康、旅游/出行、电商/零售、农业、网络安全、电力/能源、'
    '法律科技、HR/招聘、广告/创意、物流/配送、智慧城市、游戏、餐饮、体育健身、服装时尚、可再生能源、电信/5G等。'
    '案例的地域分布以中国和亚洲为主（约83%），国际案例约35+例（占比约17%），覆盖北美、欧洲、非洲和拉美地区。'
    '210案例中，成功案例203个（约96.7%），失败/争议案例7个（约3.3%）。'
)
add_paragraph(s7_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

# FIGURE 3: Principle Frequency
add_figure(
    os.path.join(FIG_DIR, 'fig3_principle_frequency.png'),
    '图3  SCOPE 36项AI业务发明原理验证频次排名（n=210案例）',
    'Fig.3  Validation Frequency Ranking of 36 AI Business Inventive Principles (n=210 cases)',
    width_inches=5.3
)

add_subheading('7.3  关键增长变化与失败案例分析')
s7_3 = (
    '从150案例到210案例的扩展过程中，几个显著变化值得关注：数字孪生（原理19）从23到34次验证（+11），是新增案例的最大贡献者，'
    '建筑/施工、考古/文化遗产、半导体制造等行业的数字孪生应用是主要增长来源。生成式设计（原理22）从6到13次验证（+7），'
    '芯片设计的AI EDA和生物制剂的AI抗体设计是新增长极。数字员工/智能体（原理14）从18到28次验证（+10），'
    '与2025-2026年Agent架构的爆发趋势高度吻合。三大支柱绝对稳固——原理2（+19）和原理7（+18）继续保持遥遥领先。'
)
add_paragraph(s7_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s7_3b = (
    '7个失败/争议案例为SCOPE提供了宝贵的"反例"数据库。其中最具代表性的是：携程AI调价助手下线事件（2026年3月），'
    '展现了S5（收入增长）× C2（可解释性）矛盾对下原理4（动态定价与匹配）的典型风险——算法定价不透明→"大数据杀熟"→'
    '信任崩塌→监管介入→被迫下线。Zillow iBuying算法的失败：AI房价预测模型过度乐观→以高于市场价买入大量房产→'
    '2.5亿美元库存减值→关闭业务，是S4（风险控制）× C4（成本）矛盾对下预测模型偏差引发重大经营风险的典型案例。'
    'Amazon AI招聘偏见事件：基于历史数据训练的招聘AI系统对女性候选人的系统性歧视，是C2和C5双重恶化参数的典型案例。'
    '失败案例的收录是SCOPE方法论的一个重要特色——它们为矩阵的禁忌标记提供了宝贵的实证依据，类似于医学领域的"不良反应报告"机制。'
)
add_paragraph(s7_3b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 8. SCOPE AGENT SYSTEM
# ============================================================
add_heading_custom('8  SCOPE Agent系统实现')

s8 = (
    '为将SCOPE方法论从纸上架构转化为可交互的工程系统，开发了SCOPE Agent Web应用，实现了五步法引导流程的自动化。'
    '系统采用前后端分离架构：前端基于Next.js 16 + React 19 + Tailwind CSS 4构建响应式Web界面；'
    '后端通过Next.js API Routes连接AI推理引擎（通义千问/DashScope）；数据持久化采用Prisma 7 + PostgreSQL'
    '（部署于阿里云RDS）；用户认证采用JWT（jose库）；整体部署于阿里云ECS（Ubuntu 22.04，Docker容器化）。'
)
add_paragraph(s8, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s8b = (
    'SCOPE Agent将五步法转化为对话式引导流程：第一步Strengthen，Agent展示6个改善参数及其定义，询问用户"最想提升什么？"，'
    '支持多选并自动判断主次。第二步Constrain，Agent展示6个恶化参数，询问用户"最担心什么？"，引导用户识别可能被忽视的约束。'
    '第三步Originate，Agent自动查询矛盾矩阵，输出推荐原理列表，并按信度等级和禁忌标记排序，每个推荐原理附带完整的五段模板信息'
    '和典型案例。第四步Path，Agent评估用户业务场景在7条进化路线上的当前阶段，规划近期（复制/自动化）→ 中期（增强/预测）→ '
    '远期（自主/创造）的三级跳路径。第五步Effect，Agent根据推荐的发明原理，从效应库中选择匹配的AI技术方案，'
    '并考虑适用条件（数据量、标注需求、计算资源、人才储备等）。SCOPE Agent已于2026年7月部署上线，'
    '公网访问地址为 http://121.40.61.131:3000。系统内置了完整的矛盾矩阵数据库、36原理完整模板库、7条进化路线描述和14项效应映射。'
)
add_paragraph(s8b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 9. DISCUSSION
# ============================================================
add_heading_custom('9  讨论')

add_subheading('9.1  方法论的创新与贡献')
s9_1 = (
    'SCOPE方法论的核心贡献在于将AI应用场景的发现从"拍脑袋"的艺术转变为"按图索骥"的科学。具体而言：'
    '（1）矛盾建模的创新——将TRIZ的技术矛盾概念迁移到业务领域，用6+6个业务场景参数替代39个工程参数。'
    '参数的"粗粒度"设计（6×6而非39×39）是有意为之——降低非技术用户的认知门槛，使业务人员而非AI专家也能参与场景建模。'
    '（2）基于实证的推荐体系——210个案例的交叉统计为矛盾矩阵的每个单元格提供了可量化的信度基础，'
    '这与大多数"AI应用场景匹配"工具依赖专家经验规则的做法有着本质区别。'
    '（3）失败案例的价值化——SCOPE主动收录并分析失败案例（7例），将其转化为矩阵中的禁忌标记，'
    '这一做法在AI应用方法论领域较为罕见，知道"什么不该做"与知道"该做什么"同等重要。'
    '（4）时间维度的引入——7条进化路线和三级跳路径规划为AI应用场景提供了动态视角。'
    '（5）工程化验证——SCOPE Agent系统的开发和部署，将方法论从纸面架构转化为可运行、可测试的工程系统。'
)
add_paragraph(s9_1, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('9.2  技术民主化集群')
s9_2 = (
    '210案例的分析揭示了一个具有理论意义的模式——SCOPE中存在一条清晰的"技术民主化集群"：'
    '原理16（能力即服务/API化）→ 原理26（能力延伸/赋能个体）→ 原理36（客户共创/赋能社群）→ '
    '原理37（元能力/赋能AI自身）。这条集群描述了一个从"将专业能力产品化"到"让每个人拥有专家级能力"到'
    '"催化社群级创造"到"AI学会如何学习"的赋能升级路径。这一发现超越了SCOPE作为"实用工具"的定位，'
    '指向了AI技术民主化的深层逻辑。'
)
add_paragraph(s9_2, font_name='宋体', font_size=Pt(10.5), space_after=Pt(4))

add_subheading('9.3  局限性与未来工作')
s9_3 = (
    'SCOPE方法论存在以下主要局限性：（1）案例规模与统计信度——210案例 vs TRIZ的4万份专利，差距为两个数量级，'
    '矩阵中仍有9个单元格处于最低信度（仅1-2案例），基于该信度的推荐应被理解为"冷启动方向性建议"而非"统计规律"。'
    '（2）案例选择偏差——案例以"成功的、可公开的、有量化指标的"为主，"静默失败"大量遗漏，'
    '可能导致方法论对风险的评估系统性偏乐观。（3）地域偏差——83%的案例来自中国和亚洲，国际案例仅占17%，'
    '方法论在全球范围内的适用性有待进一步验证。（4）时间窗口效应——所有案例来自2025-2026年大模型爆发期，'
    '当AI技术进入下一个代际，当前提炼的原理和矩阵是否仍然有效，是必须面对的不确定性。'
    '（5）参数粒度的权衡——6×6的参数设计虽然降低了使用门槛，但S1（运营效率）覆盖73%的案例，'
    '其过度聚合问题在210案例规模下依然存在。'
)
add_paragraph(s9_3, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s9_3b = (
    '未来工作方向包括：（1）案例库扩展至500+目标，重点填补低信度单元格；'
    '（2）进行"盲推"前向验证——用SCOPE在尚未被案例覆盖的新行业中独立预测AI应用场景，与实际发展情况进行对比，'
    '检验方法论的预测能力；（3）开发SCOPE的多语言版本和行业定制化版本；'
    '（4）建立开放的社区贡献机制和在线案例库。'
)
add_paragraph(s9_3b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(6))

# ============================================================
# 10. CONCLUSION
# ============================================================
add_heading_custom('10  结论')

s10 = (
    '本文提出了SCOPE——一种面向AI应用场景发现的系统性方法论。该方法论借鉴TRIZ的思想结构，但基于210个真实AI应用案例的'
    '统计分析，构建了完全独立的五层架构：业务场景参数体系（6+6参数）、AI业务发明原理库（36个原理，五大类）、矛盾矩阵'
    '（6×6，100%填满率，信度等级+禁忌标记）、AI业务进化路线（7条路线+回滚分支）、以及AI能力效应库（14项映射）。'
)
add_paragraph(s10, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s10b = (
    '210案例的统计验证表明：（1）认知自动化（63次）和预测与推演（60次）构成AI赋能业务的"双引擎"；'
    '（2）运营效率×实施复杂度（N≈55）是最强矛盾-原理组合；（3）数字孪生（+11）和生成式设计（+7）是增长最快的原理方向。'
    '基于该方法论开发的SCOPE Agent系统已部署上线，实现了五步法引导流程的工程化和可交互化。'
)
add_paragraph(s10b, font_name='宋体', font_size=Pt(10.5), space_after=Pt(2))

s10c = (
    'SCOPE的核心价值在于：将AI应用场景的发现，从依赖个人灵感的"艺术"变成了一套可操作、可验证、可演化的"科学"方法。'
    '它为企业在AI时代的战略规划提供了一种新的思维范式——不是被动等待"AI能做什么"的技术公告，而是主动用矛盾建模、原理推荐、'
    '进化路径的系统工具，绘制自己业务的AI全景图。'
)
add_paragraph(s10c, font_name='宋体', font_size=Pt(10.5), space_after=Pt(8))

# ============================================================
# REFERENCES
# ============================================================
add_heading_custom('参考文献', font_name='黑体', font_size=Pt(14), space_before=Pt(14))

references = [
    '[1] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]//Advances in Neural Information Processing Systems (NeurIPS). Virtual, 2020, 33: 1877-1901.',
    '[2] OpenAI. GPT-4 Technical Report[R]. arXiv preprint arXiv:2303.08774, 2023.',
    '[3] DeepSeek-AI. DeepSeek-V3 Technical Report[R]. arXiv preprint arXiv:2412.19437, 2024.',
    '[4] Brynjolfsson E, Mitchell T. What Can Machine Learning Do? Workforce Implications[J]. Science, 2017, 358(6370): 1530-1534.',
    '[5] Agrawal A, Gans J, Goldfarb A. Prediction Machines: The Simple Economics of Artificial Intelligence[M]. Boston: Harvard Business Review Press, 2018.',
    '[6] Ng A. AI Transformation Playbook: How to Lead Your Company to the AI Era[R]. Landing AI, 2018.',
    '[7] Iansiti M, Lakhani K R. Competing in the Age of AI: Strategy and Leadership When Algorithms and Networks Run the World[M]. Boston: Harvard Business Review Press, 2020.',
    '[8] Fountaine T, McCarthy B, Saleh T. Building the AI-Powered Organization[J]. Harvard Business Review, 2019, 97(4): 62-73.',
    '[9] Altshuller G. Creativity as an Exact Science[M]. New York: Gordon and Breach, 1984.',
    '[10] Altshuller G. The Innovation Algorithm: TRIZ, Systematic Innovation and Technical Creativity[M]. Worcester: Technical Innovation Center, 1999.',
    '[11] Altshuller G. 创新40法：TRIZ创造性解决技术问题的诀窍[M]. 黄玉霖, 范怡红, 译. 成都: 西南交通大学出版社, 2004.',
    '[12] Ilevbare I M, Probert D, Phaal R. A Review of TRIZ, and Its Benefits and Challenges in Practice[J]. Technovation, 2013, 33(2-3): 30-37.',
    '[13] Chechurin L, Borgianni Y. Understanding TRIZ Through the Review of Top Cited Publications[J]. Computers in Industry, 2016, 82: 119-134.',
    '[14] 赵敏, 史晓凌, 段海波. TRIZ入门及实践[M]. 北京: 科学出版社, 2009.',
    '[15] 檀润华. TRIZ及应用：技术创新过程与方法[M]. 北京: 高等教育出版社, 2010.',
    '[16] Barnett J, Kieslich K, Sinchai J, et al. Scenarios in Computing Research: A Systematic Review of the Use of Scenario Methods for Exploring the Future of Computing Technologies in Society[C]//Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society (AIES). 2025.',
    '[17] Negrini A. AI and Foresight: A Systematic Literature Review on Generative AI for Future Scenario Generation[D]. Milano: Politecnico di Milano, 2024.',
    '[18] Wang B, Li Y, Lei W. PRISE: A Framework for AI Product Incubation from Concept to Implementation[J]. IEEE Access, 2025, 13: 119740-119756.',
    '[19] Dsouza R, et al. Innovation Process for AI-Enabled Products and Services (IPAPS)[C]//Proceedings of the Pacific Asia Conference on Information Systems (PACIS). 2025.',
    '[20] Walter T. Finding Value in AI: A Framework for Practical Adoption[R]. Merkle, 2025.',
    '[21] Scheichl P. Artificial Intelligence for Autonomous Innovation: Extending the Double Diamond Model with AI Capabilities[D]. Steyr: FH Oberösterreich, 2024.',
    '[22] 赵健, 等. 人工智能大模型及其应用综述[J]. Visual Intelligence, 2024.',
    '[23] 大模型行业应用设计与实践[J]. 计算机与网络, 2024(4).',
    '[24] 罗毅晗. SCOPE — AI生成应用场景方法论（V4.1.1）[R]. 2026.',
    '[25] Salamatov Y. TRIZ: The Right Solution at the Right Time[M]. Netherlands: Insytec, 1999.',
    '[26] Savransky S D. Engineering of Creativity: Introduction to TRIZ Methodology of Inventive Problem Solving[M]. Boca Raton: CRC Press, 2000.',
    '[27] Mann D. Hands-On Systematic Innovation[M]. 2nd ed. Clevedon: IFR Press, 2007.',
    '[28] 孙永伟, 谢尔盖·伊克万科. TRIZ：打开创新之门的金钥匙[M]. 北京: 科学出版社, 2014.',
    '[29] Radziwill N M. Disconnected Data: The Challenge of AI Integration[J]. Quality Management Journal, 2019, 26(1): 7-11.',
    '[30] Verganti R, Vendraminelli L, Iansiti M. Innovation and Design in the Age of Artificial Intelligence[J]. Journal of Product Innovation Management, 2020, 37(3): 212-227.',
    '[31] Zhou J, et al. A Survey on the Applications of Generative AI in Automated Driving Systems Test Scenario Generation Methods[R]. arXiv preprint arXiv:2512.15422, 2025.',
    '[32] Bonaccorsi A, et al. Artificial Intelligence and Innovation: A Systematic Literature Review and Future Research Directions[J]. Technological Forecasting and Social Change, 2024, 198: 122992.',
]

for ref in references:
    add_paragraph(ref, font_name='Times New Roman', font_size=Pt(9),
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(0),
                  space_before=Pt(0), space_after=Pt(1))

# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
print(f"\nWord document saved to: {DOC_PATH}")
print("Done!")
